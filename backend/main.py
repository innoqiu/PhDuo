# from fastapi import FastAPI
# from fastapi.middleware.cors import CORSMiddleware

# from app.database import create_tables
# from app.api.endpoints import router

# # Create FastAPI app
# app = FastAPI(title="Academic Matching Assistant", version="1.0.0")

# # Configure CORS
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["*"],
#     allow_credentials=True,
#     allow_methods=["*"],
#     allow_headers=["*"],
# )

# # Include API routes  
# app.include_router(router, tags=["analysis"])

# # Create database tables on startup
# @app.on_event("startup")
# async def startup_event():
#     create_tables()

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=7878)


import os
import json
import asyncio
from typing import Optional, Callable, Any, Tuple
from urllib.parse import urlparse
import re
from pathlib import Path
import hashlib
from datetime import datetime
import time
import random
import ipaddress
from openai import RateLimitError, APIError, APIConnectionError, APITimeoutError

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import database_operations as db_ops
from sqlalchemy import create_engine, Column, String, Text, Integer, JSON
from sqlalchemy.orm import sessionmaker, declarative_base
from openai import AsyncOpenAI
from json_repair import repair_json
from app.services.crawl4ai_service import crawl_professor_website
import pypdf
from docx import Document
from dotenv import load_dotenv

load_dotenv()
# --- 配置 ---
ZEABUR_API_KEY = os.getenv("ZEABUR_API_KEY")
ZEABUR_BASE_URL = os.getenv("ZEABUR_BASE_URL", "https://hnd1.aihub.zeabur.ai/v1")
DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./professors.db")

# 初始化客户端
aclient = AsyncOpenAI(api_key=ZEABUR_API_KEY, base_url=ZEABUR_BASE_URL)

# --- Retry Logic with Exponential Backoff ---

async def retry_llm_call(
    func: Callable,
    max_retries: int = 5,
    initial_delay: float = 1.0,
    max_delay: float = 60.0,
    exponential_base: float = 2.0,
    jitter: bool = True
) -> Any:
    """
    Retry an async LLM call with exponential backoff.
    Handles rate limits (429), server errors (5xx), and connection errors.
    
    Args:
        func: Async function to retry (should be a coroutine function, not a coroutine)
        max_retries: Maximum number of retry attempts
        initial_delay: Initial delay in seconds before first retry
        max_delay: Maximum delay in seconds between retries
        exponential_base: Base for exponential backoff calculation
        jitter: Whether to add random jitter to delay
    
    Returns:
        Result of the function call
    
    Raises:
        Last exception if all retries are exhausted
    """
    last_exception = None
    
    for attempt in range(max_retries + 1):
        try:
            # Call the function
            result = await func()
            return result
            
        except (RateLimitError, APIError, APIConnectionError, APITimeoutError) as e:
            last_exception = e
            
            # Check if it's a rate limit error (HTTP 429) or other retryable error
            is_rate_limit = isinstance(e, RateLimitError)
            status_code = None
            
            # Try to extract status code from exception
            if hasattr(e, 'status_code'):
                status_code = e.status_code
                is_rate_limit = (status_code == 429) or is_rate_limit
            elif hasattr(e, 'response'):
                if hasattr(e.response, 'status_code'):
                    status_code = e.response.status_code
                    is_rate_limit = (status_code == 429) or is_rate_limit
                elif hasattr(e.response, 'status'):
                    status_code = e.response.status
                    is_rate_limit = (status_code == 429) or is_rate_limit
            
            # Determine if error is retryable
            retryable_status_codes = [429, 500, 502, 503, 504]
            is_retryable = is_rate_limit or (status_code in retryable_status_codes) or isinstance(e, (APIConnectionError, APITimeoutError))
            
            # Don't retry on non-retryable errors
            if not is_retryable and status_code and status_code not in retryable_status_codes:
                print(f"❌ Non-retryable error (status {status_code}): {str(e)}")
                raise
            
            # If this was the last attempt, raise the exception
            if attempt >= max_retries:
                print(f"❌ Max retries ({max_retries}) exceeded. Last error: {str(e)}")
                raise
            
            # Calculate delay with exponential backoff
            delay = min(initial_delay * (exponential_base ** attempt), max_delay)
            
            # Add jitter to prevent thundering herd
            if jitter:
                jitter_amount = random.uniform(0, delay * 0.1)  # Up to 10% jitter
                delay += jitter_amount
            
            # Log retry attempt
            if is_rate_limit:
                error_type = "Rate Limit (429)"
            elif status_code:
                error_type = f"API Error ({status_code})"
            elif isinstance(e, APIConnectionError):
                error_type = "Connection Error"
            elif isinstance(e, APITimeoutError):
                error_type = "Timeout Error"
            else:
                error_type = "API Error"
            
            print(f"⚠️  {error_type} encountered (attempt {attempt + 1}/{max_retries + 1}). Retrying in {delay:.2f} seconds...")
            print(f"   Error: {str(e)}")
            
            # Wait before retrying
            await asyncio.sleep(delay)
            
        except Exception as e:
            # For non-retryable errors (not OpenAI API errors), raise immediately
            if not isinstance(e, (RateLimitError, APIError, APIConnectionError, APITimeoutError)):
                raise
            last_exception = e
    
    # If we get here, all retries were exhausted
    if last_exception:
        raise last_exception
    else:
        raise Exception("Retry logic exhausted without result")

# --- 数据库设置 (SQLite) ---
engine = create_engine(DATABASE_URL, connect_args={"check_same_thread": False})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()

class Professor(Base):
    __tablename__ = "professors"
    id = Column(Integer, primary_key=True, index=True)
    root_url = Column(String, unique=True, index=True) # 用于索引
    name = Column(String)
    university = Column(String)
    profile_data = Column(JSON) # 存储 LLM 提取的结构化 Profile

class MatchReport(Base):
    __tablename__ = "match_reports"
    id = Column(Integer, primary_key=True, index=True)
    professor_name = Column(String, index=True)
    student_name = Column(String, index=True)
    professor_url = Column(String)
    created_at = Column(String)  # ISO format timestamp
    refined_report = Column(JSON)  # The structured refined report data
    overall_score = Column(String)  # Store as string for easy access

class ProfessorAnalysis(Base):
    __tablename__ = "professor_analyses"
    id = Column(Integer, primary_key=True, index=True)
    professor_url = Column(String, unique=True, index=True)  # Unique per professor
    professor_name = Column(String, index=True)
    analysis_report = Column(Text)  # The professor analysis report text
    created_at = Column(String)  # ISO format timestamp
    profile_hash = Column(String)  # Hash of profile_data to detect changes

Base.metadata.create_all(bind=engine)

app = FastAPI()

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 辅助函数 ---

def load_prompt(filename: str, **kwargs) -> str:
    """
    Load a prompt template from the prompts directory and format it with provided variables
    
    Args:
        filename: Name of the prompt file (e.g., 'cv_extraction.txt')
        **kwargs: Variables to format into the prompt template
        
    Returns:
        Formatted prompt string
    """
    prompts_dir = Path(__file__).parent / "prompts"
    prompt_file = prompts_dir / filename
    
    if not prompt_file.exists():
        raise FileNotFoundError(f"Prompt file not found: {prompt_file}")
    
    with open(prompt_file, 'r', encoding='utf-8') as f:
        template = f.read()
    
    # Format the template with provided variables
    return template.format(**kwargs)

def url_to_filename(url: str) -> str:
    """
    将 URL 转换为安全的文件名
    使用 URL hash 来避免文件名冲突和特殊字符问题
    """
    # 创建 URL 的 hash
    url_hash = hashlib.md5(url.encode('utf-8')).hexdigest()[:8]
    # 提取域名作为可读前缀
    parsed = urlparse(url)
    domain = parsed.netloc.replace('.', '_').replace(':', '_')
    # 组合成安全的文件名
    return f"{domain}_{url_hash}.txt"

def ensure_cache_directory() -> Path:
    """
    确保缓存目录存在
    """
    cache_dir = Path("crawl_cache")
    cache_dir.mkdir(exist_ok=True)
    return cache_dir

def ensure_match_log_directory() -> Path:
    """
    确保匹配日志目录存在
    """
    log_dir = Path("match_log")
    log_dir.mkdir(exist_ok=True)
    return log_dir

def ensure_metadata_directory() -> Path:
    """
    确保元数据目录存在
    """
    metadata_dir = Path("metadata")
    metadata_dir.mkdir(exist_ok=True)
    return metadata_dir

def ensure_cv_cache_directory() -> Path:
    """
    确保CV缓存目录存在
    """
    cv_cache_dir = Path("cv_cache")
    cv_cache_dir.mkdir(exist_ok=True)
    return cv_cache_dir

def compute_file_hash(file_content: bytes) -> str:
    """
    计算文件内容的SHA256哈希值，用于唯一标识文件
    """
    return hashlib.sha256(file_content).hexdigest()

def get_cv_cache_path(file_hash: str) -> Path:
    """
    获取CV缓存文件路径
    """
    cv_cache_dir = ensure_cv_cache_directory()
    return cv_cache_dir / f"{file_hash}.json"

def save_cv_cache(file_hash: str, raw_text: str, fixed_text: dict) -> None:
    """
    保存CV提取结果到缓存
    """
    try:
        cache_file = get_cv_cache_path(file_hash)
        cache_data = {
            "file_hash": file_hash,
            "raw_text": raw_text,
            "fixed_text": fixed_text,
            "cached_at": datetime.now().isoformat()
        }
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=2)
        print(f"CV cache saved: {cache_file}")
    except Exception as e:
        print(f"Error saving CV cache: {str(e)}")

def get_cv_cache(file_hash: str) -> Optional[dict]:
    """
    从缓存获取CV提取结果
    返回包含 raw_text 和 fixed_text 的字典，如果不存在则返回 None
    """
    try:
        cache_file = get_cv_cache_path(file_hash)
        if not cache_file.exists():
            return None
        
        with open(cache_file, 'r', encoding='utf-8') as f:
            cache_data = json.load(f)
        
        # 验证哈希值匹配
        if cache_data.get("file_hash") != file_hash:
            print(f"Warning: Cache file hash mismatch for {file_hash}")
            return None
        
        print(f"CV cache hit: {cache_file}")
        return cache_data
    except Exception as e:
        print(f"Error reading CV cache: {str(e)}")
        return None

def sanitize_filename(name: str) -> str:
    """
    清理文件名，移除不允许的字符
    """
    # 移除或替换不允许的文件名字符
    name = re.sub(r'[<>:"/\\|?*]', '_', name)
    # 移除多余的空格和下划线
    name = re.sub(r'\s+', '_', name.strip())
    name = re.sub(r'_+', '_', name)
    # 限制长度
    if len(name) > 100:
        name = name[:100]
    return name if name else "Unknown"

def save_match_log(complete_text: str, professor_profile: dict, student_profile: dict) -> str:
    """
    保存匹配分析结果到日志文件
    文件名格式: timestamp_professorname_studentname.txt
    """
    try:
        log_dir = ensure_match_log_directory()
        
        # 提取教授姓名
        professor_name = "Unknown"
        if isinstance(professor_profile, dict):
            identity = professor_profile.get("identity", {})
            if isinstance(identity, dict):
                professor_name = identity.get("name", "Unknown")
        
        # 提取学生姓名
        student_name = "Unknown"
        if isinstance(student_profile, dict):
            identity = student_profile.get("identity", {})
            if isinstance(identity, dict):
                student_name = identity.get("full_name", "Unknown")
        
        # 清理文件名
        professor_name_clean = sanitize_filename(str(professor_name))
        student_name_clean = sanitize_filename(str(student_name))
        
        # 生成时间戳 (格式: YYYYMMDD_HHMMSS)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 生成文件名
        filename = f"{timestamp}_{professor_name_clean}_{student_name_clean}.txt"
        log_file = log_dir / filename
        
        # 写入文件
        with open(log_file, 'w', encoding='utf-8') as f:
            f.write(f"# Match Analysis Log\n")
            f.write(f"# Generated: {datetime.now().isoformat()}\n")
            f.write(f"# Professor: {professor_name}\n")
            f.write(f"# Student: {student_name}\n")
            f.write(f"# {'='*60}\n\n")
            f.write(complete_text)
        
        print(f"Match log saved to: {log_file}")
        return str(log_file)
        
    except Exception as e:
        print(f"Error saving match log: {str(e)}")
        return ""

def save_refined_report(refined_report_json: dict, professor_profile: dict, student_profile: dict) -> str:
    """
    保存精炼报告到元数据文件夹
    文件名格式: timestamp_professorname_studentname.txt
    """
    try:
        metadata_dir = ensure_metadata_directory()
        
        # 提取教授姓名
        professor_name = "Unknown"
        if isinstance(professor_profile, dict):
            identity = professor_profile.get("identity", {})
            if isinstance(identity, dict):
                professor_name = identity.get("name", "Unknown")
        
        # 提取学生姓名
        student_name = "Unknown"
        if isinstance(student_profile, dict):
            identity = student_profile.get("identity", {})
            if isinstance(identity, dict):
                student_name = identity.get("full_name", "Unknown")
        
        # 清理文件名
        professor_name_clean = sanitize_filename(str(professor_name))
        student_name_clean = sanitize_filename(str(student_name))
        
        # 生成时间戳 (格式: YYYYMMDD_HHMMSS)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 生成文件名
        filename = f"{timestamp}_{professor_name_clean}_{student_name_clean}.txt"
        metadata_file = metadata_dir / filename
        
        # 写入文件 (保存为格式化的JSON文本)
        with open(metadata_file, 'w', encoding='utf-8') as f:
            f.write(json.dumps(refined_report_json, ensure_ascii=False, indent=2))
        
        print(f"Refined report saved to: {metadata_file}")
        return str(metadata_file)
        
    except Exception as e:
        print(f"Error saving refined report: {str(e)}")
        return ""

async def get_professor_analysis_cache(professor_url: str, professor_profile: dict) -> Optional[str]:
    """
    从数据库获取缓存的教授分析报告（异步版本）
    如果教授资料发生变化（通过profile_hash检测），返回None
    返回分析报告文本，如果不存在或已过期则返回None
    """
    def _get_cache():
        try:
            db = SessionLocal()
            
            # 计算当前profile的hash
            profile_json = json.dumps(professor_profile, sort_keys=True, ensure_ascii=False)
            current_profile_hash = hashlib.sha256(profile_json.encode('utf-8')).hexdigest()
            
            # 查找缓存
            cached_analysis = db.query(ProfessorAnalysis).filter(
                ProfessorAnalysis.professor_url == professor_url
            ).first()
            
            if cached_analysis:
                # 检查profile是否发生变化
                if cached_analysis.profile_hash == current_profile_hash:
                    print(f"Professor analysis cache hit for {professor_url}")
                    analysis_report = cached_analysis.analysis_report
                    db.close()
                    return analysis_report
                else:
                    print(f"Professor profile changed, invalidating cache for {professor_url}")
                    # Profile changed, delete old cache
                    db.delete(cached_analysis)
                    db.commit()
            
            db.close()
            return None
        except Exception as e:
            print(f"Error reading professor analysis cache: {str(e)}")
            if 'db' in locals():
                db.close()
            return None
    
    return await run_db_operation(_get_cache)

async def save_professor_analysis_to_db(professor_url: str, professor_name: str, analysis_report: str, professor_profile: dict) -> None:
    """
    保存教授分析报告到数据库（异步版本）
    """
    def _save_cache():
        try:
            db = SessionLocal()
            
            # 计算profile的hash
            profile_json = json.dumps(professor_profile, sort_keys=True, ensure_ascii=False)
            profile_hash = hashlib.sha256(profile_json.encode('utf-8')).hexdigest()
            
            # 检查是否已存在
            existing = db.query(ProfessorAnalysis).filter(
                ProfessorAnalysis.professor_url == professor_url
            ).first()
            
            if existing:
                # 更新现有记录
                existing.analysis_report = analysis_report
                existing.profile_hash = profile_hash
                existing.created_at = datetime.now().isoformat()
                print(f"Updated professor analysis cache for {professor_url}")
            else:
                # 创建新记录
                new_analysis = ProfessorAnalysis(
                    professor_url=professor_url,
                    professor_name=professor_name,
                    analysis_report=analysis_report,
                    profile_hash=profile_hash,
                    created_at=datetime.now().isoformat()
                )
                db.add(new_analysis)
                print(f"Saved new professor analysis cache for {professor_url}")
            
            db.commit()
            db.close()
        except Exception as e:
            print(f"Error saving professor analysis cache: {str(e)}")
            if 'db' in locals():
                db.rollback()
                db.close()
    
    await run_db_operation(_save_cache)

async def save_match_report_to_db(refined_report_json: dict, professor_profile: dict, student_profile: dict, professor_url: str) -> Optional[int]:
    """
    保存匹配报告到数据库（异步版本）
    返回报告ID，如果失败返回None
    """
    def _save_report():
        try:
            db = SessionLocal()
            
            # 提取教授姓名
            professor_name = "Unknown"
            if isinstance(professor_profile, dict):
                identity = professor_profile.get("identity", {})
                if isinstance(identity, dict):
                    professor_name = identity.get("name", "Unknown")
            
            # 提取学生姓名
            student_name = "Unknown"
            if isinstance(student_profile, dict):
                identity = student_profile.get("identity", {})
                if isinstance(identity, dict):
                    student_name = identity.get("full_name", "Unknown")
            
            # 获取整体分数
            overall_score = str(refined_report_json.get("meta", {}).get("overallScore", "N/A"))
            created_at = refined_report_json.get("meta", {}).get("generated", datetime.now().isoformat())
            
            # 创建新的匹配报告记录
            new_report = MatchReport(
                professor_name=professor_name,
                student_name=student_name,
                professor_url=professor_url,
                created_at=created_at,
                refined_report=refined_report_json,
                overall_score=overall_score
            )
            
            db.add(new_report)
            db.commit()
            db.refresh(new_report)
            
            report_id = new_report.id
            db.close()
            
            print(f"Match report saved to database with ID: {report_id}")
            return report_id
            
        except Exception as e:
            print(f"Error saving match report to database: {str(e)}")
            import traceback
            traceback.print_exc()
            if 'db' in locals():
                db.rollback()
                db.close()
            return None
    
    return await run_db_operation(_save_report)

def get_cached_content(root_url: str) -> str | None:
    """
    检查是否有已缓存的爬取内容
    """
    try:
        cache_dir = ensure_cache_directory()
        filename = url_to_filename(root_url)
        cache_file = cache_dir / filename
        
        if cache_file.exists():
            with open(cache_file, 'r', encoding='utf-8') as f:
                content = f.read()
                print(f"Cache hit: loaded {len(content)} characters from {cache_file}")
                return content
        return None
    except Exception as e:
        print(f"Error reading cache for {root_url}: {str(e)}")
        return None

def save_cached_content(root_url: str, content: str) -> None:
    """
    保存爬取内容到本地缓存文件
    """
    try:
        cache_dir = ensure_cache_directory()
        filename = url_to_filename(root_url)
        cache_file = cache_dir / filename
        
        with open(cache_file, 'w', encoding='utf-8') as f:
            # 添加元数据头部
            from datetime import datetime
            f.write(f"# Crawled URL: {root_url}\n")
            f.write(f"# Cached at: {datetime.now().isoformat()}\n")
            f.write(f"# Content length: {len(content)} characters\n")
            f.write("# " + "="*50 + "\n\n")
            f.write(content)
            
        print(f"Cached {len(content)} characters to {cache_file}")
    except Exception as e:
        print(f"Error saving cache for {root_url}: {str(e)}")

def list_cached_urls() -> list[str]:
    """
    列出所有已缓存的 URL
    """
    try:
        cache_dir = ensure_cache_directory()
        cached_urls = []
        
        for cache_file in cache_dir.glob("*.txt"):
            try:
                with open(cache_file, 'r', encoding='utf-8') as f:
                    first_line = f.readline()
                    if first_line.startswith("# Crawled URL: "):
                        url = first_line.replace("# Crawled URL: ", "").strip()
                        cached_urls.append(url)
            except Exception:
                continue
                
        return cached_urls
    except Exception as e:
        print(f"Error listing cached URLs: {str(e)}")
        return []

def clear_cache(url: str = None) -> bool:
    """
    清除缓存（单个 URL 或全部）
    """
    try:
        cache_dir = ensure_cache_directory()
        
        if url:
            # 清除单个 URL 的缓存
            filename = url_to_filename(url)
            cache_file = cache_dir / filename
            if cache_file.exists():
                cache_file.unlink()
                print(f"Cleared cache for {url}")
                return True
            else:
                print(f"No cache found for {url}")
                return False
        else:
            # 清除所有缓存
            count = 0
            for cache_file in cache_dir.glob("*.txt"):
                cache_file.unlink()
                count += 1
            print(f"Cleared {count} cache files")
            return True
            
    except Exception as e:
        print(f"Error clearing cache: {str(e)}")
        return False

def validate_url(url: str) -> Tuple[bool, str]:
    """
    Validate that URL is a public HTTPS URL.
    
    Args:
        url: URL string to validate
        
    Returns:
        Tuple of (is_valid, error_message)
        If valid: (True, "")
        If invalid: (False, error_message)
    """
    try:
        parsed = urlparse(url)
        
        # Check scheme is HTTPS
        if parsed.scheme != "https":
            return False, f"URL must use HTTPS scheme. Found: {parsed.scheme}"
        
        # Check that netloc exists
        if not parsed.netloc:
            return False, "URL must have a valid domain (netloc)"
        
        # Extract hostname (remove port if present)
        hostname = parsed.netloc.split(':')[0]
        
        # Check for localhost and local domains
        local_domains = [
            "localhost",
            "127.0.0.1",
            "0.0.0.0",
            "::1",
            "local",
        ]
        
        hostname_lower = hostname.lower()
        if hostname_lower in local_domains:
            return False, f"URL cannot point to localhost or local domains. Found: {hostname}"
        
        # Check if hostname is an IP address
        try:
            ip = ipaddress.ip_address(hostname)
            
            # Check if it's a private/internal IP
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                return False, f"URL cannot use private/internal IP addresses. Found: {hostname}"
            
            # Allow public IPs (though domain names are preferred)
            # This is acceptable for public IPs
            return True, ""
            
        except ValueError:
            # Not an IP address, check if it's a valid domain
            
            # Check for local/internal domain patterns
            internal_domain_patterns = [
                '.local',
                '.localhost',
                '.internal',
                '.corp',
                '.lan',
                '.home',
            ]
            if any(hostname_lower.endswith(pattern) for pattern in internal_domain_patterns):
                return False, f"URL cannot use local/internal domains. Found: {hostname}"
            
            # Check for invalid characters in hostname
            if not re.match(r'^[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?(\.[a-zA-Z0-9]([a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?)*$', hostname):
                return False, f"URL contains invalid characters in hostname. Found: {hostname}"
            
            # Check that it has at least one dot (basic domain validation)
            if '.' not in hostname:
                return False, f"URL must be a valid domain name. Found: {hostname}"
            
            # Basic TLD check (must have at least 2 parts: domain.tld)
            parts = hostname.split('.')
            if len(parts) < 2:
                return False, f"URL must have a valid domain and TLD. Found: {hostname}"
            
            # All checks passed
            return True, ""
            
    except Exception as e:
        return False, f"Invalid URL format: {str(e)}"

def clean_url(url: str) -> str:
    """
    将 https://imyueli.github.io/projects.html 整理为 https://imyueli.github.io
    """
    parsed = urlparse(url)
    return f"{parsed.scheme}://{parsed.netloc}"

def detect_file_type(filename: str, file_content: bytes) -> str:
    """
    Detect file type based on filename extension and file content
    
    Returns:
        'pdf', 'docx', or 'doc'
    """
    filename_lower = filename.lower()
    
    # Check by extension first
    if filename_lower.endswith('.pdf'):
        return 'pdf'
    elif filename_lower.endswith('.docx'):
        return 'docx'
    elif filename_lower.endswith('.doc'):
        return 'doc'
    
    # Check by magic bytes (file signature)
    if file_content.startswith(b'%PDF'):
        return 'pdf'
    elif file_content.startswith(b'PK\x03\x04'):  # DOCX is a ZIP file
        # Check if it's actually a DOCX by looking for word/ in the ZIP
        try:
            import zipfile
            import io
            zip_file = zipfile.ZipFile(io.BytesIO(file_content))
            if 'word/' in zip_file.namelist():
                return 'docx'
        except:
            pass
    
    # Default to PDF for backward compatibility
    return 'pdf'

def extract_text_from_pdf(file_content: bytes) -> str:
    """简单的 PDF 文本提取"""
    import io
    pdf_reader = pypdf.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    print(f"Extracted {len(text)} characters from PDF")
    return text

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    import io
    doc = Document(io.BytesIO(file_content))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    
    print(f"Extracted {len(text)} characters from DOCX")
    return text.strip()

def extract_text_from_doc(file_content: bytes) -> str:
    """
    Extract text from .doc file (older Word format)
    Note: .doc files are binary and require special handling
    For now, we'll try to use python-docx which may not work for .doc
    If that fails, we'll raise an error suggesting conversion to .docx
    """
    # .doc files are OLE2 format and harder to parse
    # Try using antiword or similar tools, or suggest conversion
    try:
        # Try to read as docx (won't work for .doc but worth trying)
        return extract_text_from_docx(file_content)
    except:
        raise ValueError(
            "Direct .doc file parsing is not supported. "
            "Please convert your .doc file to .docx format or use PDF format."
        )

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from file based on file type (PDF, DOCX, or DOC)
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        
    Returns:
        Extracted text as string
    """
    file_type = detect_file_type(filename, file_content)
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_type == 'docx':
        return extract_text_from_docx(file_content)
    elif file_type == 'doc':
        return extract_text_from_doc(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

async def fix_pdf_with_LayoutLM(textx: str) -> str:
    """
    使用 LayoutLM 修复 PDF 文本
    """

    output_schema_template = """
                {
            "identity": {
                "full_name": null,
                "email": null,
                "phone": null,
                "location": null,
                "website": null,
                "linkedin": null,
                "github": null
            },
            "professional_summary": null,
            "education": [
                {
                "institution": null,
                "degree": null,
                "field_of_study": null,
                "start_date": null,
                "end_date": null,
                "location": null,
                "notes": null,
                "gpa_or_academic_standing": null
                }
            ],
            "work_experience": [
                {
                "company": null,
                "role": null,
                "start_date": null,
                "end_date": null,
                "location": null,
                "responsibilities": [],
                "technologies": []
                }
            ],
            "skills": {
                "programming_languages": [],
                "frameworks_and_libraries": [],
                "tools_and_platforms": [],
                "design_and_creative_tools": [],
                "other_skills": []
            },
            "projects": [
                {
                "name": null,
                "description": null,
                "role": null,
                "technologies": [],
                "year": null
                }
            ],
            "publications_or_talks": [],
            "certifications": [],
            "languages": [],
            "awards": [],
            "meta": {
                "confidence_notes": null
            }
            }
            """
    prompt_2 = load_prompt("cv_extraction.txt", output_schema_template=output_schema_template, textx=textx) 

    print("\n" + "="*60)
    print("🤖 LLM CALL #1: CV Processing (fix_pdf_with_LayoutLM)")
    print(f"   Model: gpt-5")
    print(f"   Purpose: Extract and structure CV data")
    print(f"   Input length: {len(textx)} characters")
    print("="*60)
    
    async def make_cv_call():
        return await aclient.chat.completions.create(
            model="gpt-5", # 或 zeabur
            messages=[{"role": "user", "content": prompt_2}],
            response_format={"type": "json_object"}
        )
    
    response = await retry_llm_call(make_cv_call)
    try:
        profile_json_str = response.choices[0].message.content
        #脏输出问题
        print(f"\nLLM response received for student profile extraction.")
        json_candidate = extract_json_by_brace_balance(profile_json_str)
        repaired = repair_json(json_candidate)
        profile_data = json.loads(repaired)
        print("Got Student profile data")
    except Exception as e:
        raise Exception(f"LLM response parsing failed: {str(e)}")  
    return profile_data


def extract_json_by_brace_balance(text: str) -> str:
    start = None
    depth = 0

    for i, ch in enumerate(text):
        if ch == '{':
            if depth == 0:
                start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and start is not None:
                return text[start:i+1]

    raise ValueError("No complete JSON object found")

async def get_professor_profile(root_url: str):
    """
    检查数据库，如果不存在则爬取并生成 Profile 和 Analysis
    返回 (profile_data, analysis_report) 元组
    """
    # Check database (non-blocking)
    def check_db():
        db = SessionLocal()
        try:
            return db.query(Professor).filter(Professor.root_url == root_url).first()
        finally:
            db.close()
    
    prof = await run_db_operation(check_db)
    
    profile_data = None
    analysis_report = None
    
    if prof:
        print(f"Database hit for {root_url}")
        profile_data = prof.profile_data
        
        # Check cache for analysis report
        cached_analysis = await get_professor_analysis_cache(root_url, profile_data)
        if cached_analysis:
            print("\n" + "="*60)
            print("✅ CACHE HIT: Professor Analysis")
            print(f"   Professor: {profile_data.get('identity', {}).get('name', 'Unknown')}")
            print(f"   URL: {root_url}")
            print("="*60)
            return profile_data, cached_analysis
        
        # If profile exists but analysis not cached, generate analysis from existing profile
        print(f"Profile exists but analysis not cached. Generating analysis from existing profile...")
        # Generate analysis from existing profile (Call #3 equivalent, but using existing profile)
        system_prompt = load_prompt("lab_analysis_system.txt")
        user_prompt = load_prompt(
            "lab_analysis_user.txt",
            fixed_text="[Analyzing professor profile only]",
            professor_profile=json.dumps(profile_data, ensure_ascii=False)
        )
        
        print("\n" + "="*60)
        print("🤖 LLM CALL #2+3: Professor Analysis (from existing profile)")
        print(f"   Model: gpt-5")
        print(f"   Purpose: Analyze existing professor profile")
        print(f"   Professor: {profile_data.get('identity', {}).get('name', 'Unknown')}")
        print("="*60)
        
        async def make_professor_analysis_call():
            return await aclient.chat.completions.create(
                model="gpt-5",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                stream=False
            )
        
        response = await retry_llm_call(make_professor_analysis_call)
        
        analysis_report = response.choices[0].message.content
        professor_name = profile_data.get("identity", {}).get("name", "Unknown")
        await save_professor_analysis_to_db(root_url, professor_name, analysis_report, profile_data)
        
        return profile_data, analysis_report

    # Professor not in database - do combined extraction + analysis
    print(f"Crawling new data for {root_url}")
    
    # 2. 首先检查本地缓存
    cached_content = get_cached_content(root_url)
    if cached_content:
        web_content = cached_content
        print(f"Using cached content: {len(web_content)} characters")
    else:
        # 3. 如果没有缓存，则爬取数据
        try:
            print(f"Cache miss - crawling {root_url}...")
            web_content = await crawl_professor_website(root_url)
            print(f"Crawled {len(web_content)} characters from {root_url}")
            
            # 4. 保存到本地缓存
            save_cached_content(root_url, web_content)
            
        except Exception as e:
            raise Exception(f"Crawling failed: {str(e)}")

    # 5. Combined LLM Call: Extract Profile AND Generate Analysis
    output_schema_template = """
    Return the data in this exact JSON structure. If a field is not found, return `null`.
    {
    "identity": {
        "name": "string",
        "title": "string",
        "affiliation": {
        "university": "string",
        "department": "string",
        "lab_name": "string"
        },
        "explicit_bio_data": {
        "age_or_years_active": "string",
        "background_log": ["list", "of", "strings"]
        }
    },
    "research_signals": {
        "research_interests" : "string",
        "application_or_problem_domains": ["string"],
        "venues_found": ["list", "of", "venue_acronyms"],
        "research_items": [
        {
            "title": "string",
            "type": "string", 
            "authors_raw_string": "string",
            "venue_or_journal": "string",
            "abstract_or_description": "string",
            "associated_urls": ["list", "of", "urls"]
        }
        ]
    },
    "team_raw": {
        "member_list": [
        {
            "name": "string",
            "role": "string",
            "start_year": "string",
            "involved_project_count": "string"
        }
        ],
        "sub_mentoring_signals": ["list", "of", "strings"]
    },
    "stability_signals": {
        "alumni_list": [
        {
            "name": "string",
            "current_position_raw": "string"
        }
        ],
        "ghosting_signals_from_news": ["list", "of", "names"],
        "freshness_indicators": {
        "footer_last_updated": "string",
        "latest_news_date": "string",
        "latest_publication_date": "string"
        },
        "funding_sources": ["list", "of", "strings"]
    },
    "recruitment_signals": {
        "hiring_status_raw": "string",
        "application_protocol": "string",
        "required_materials": ["list", "of", "strings"],
        "target_audience": "string"
    }
    }
    """
    combined_prompt = load_prompt("professor_extraction_and_analysis.txt", output_schema_template=output_schema_template, web_content=web_content) 

    print("\n" + "="*60)
    print("🤖 LLM CALL #2+3: Combined Professor Profile Extraction & Analysis")
    print(f"   Model: gpt-5")
    print(f"   Purpose: Extract structured profile AND generate analysis in one call")
    print(f"   URL: {root_url}")
    print(f"   Web content length: {len(web_content)} characters")
    print("="*60)
    
    async def make_combined_professor_call():
        return await aclient.chat.completions.create(
            model="gpt-5",
            messages=[{"role": "user", "content": combined_prompt}],
            response_format={"type": "json_object"}
        )
    
    response = await retry_llm_call(make_combined_professor_call)
    try:
        combined_json_str = response.choices[0].message.content
        print(f"\nLLM response received for combined extraction+analysis. length: {len(combined_json_str)}")
        json_candidate = extract_json_by_brace_balance(combined_json_str)
        repaired = repair_json(json_candidate)
        combined_data = json.loads(repaired)
        
        # Extract profile and analysis from combined response
        profile_data = combined_data.get("profile")
        analysis_report = combined_data.get("analysis_report")
        
        if not profile_data:
            raise Exception("Missing 'profile' field in LLM response")
        if not analysis_report:
            raise Exception("Missing 'analysis_report' field in LLM response")
            
    except Exception as e:
        raise Exception(f"LLM response parsing failed: {str(e)}")

    
    # 存入数据库 (non-blocking)
    def save_to_db():
        db = SessionLocal()
        try:
            new_prof = Professor(
                root_url=root_url,
                name=profile_data.get("identity", {}).get("name", "Unknown"),
                university=profile_data.get("identity", {}).get("affiliation", {}).get("university", "Unknown"),
                profile_data=profile_data
            )
            db.add(new_prof)
            db.commit()
            db.refresh(new_prof)
        finally:
            db.close()
    
    await run_db_operation(save_to_db)
    
    # Save analysis to cache
    professor_name = profile_data.get("identity", {}).get("name", "Unknown")
    await save_professor_analysis_to_db(root_url, professor_name, analysis_report, profile_data)
    
    return profile_data, analysis_report

# --- API 接口 ---

@app.post("/analyze")
async def analyze_match(
    cv: UploadFile = File(...),
    url: str = Form(...)
):
    async def event_generator():
        try:
            # Validate URL before processing
            is_valid, error_message = validate_url(url)
            if not is_valid:
                error_msg = f"Invalid URL: {error_message}. Please provide a valid public HTTPS URL."
                yield f"data: {json.dumps({'status': 'ERROR', 'error': error_msg})}\n\n"
                return
            
            # Initial setup
            root_url = clean_url(url)
            cv_content = await cv.read()
            
            # Validate file size (5MB limit)
            MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB in bytes
            file_size = len(cv_content)
            if file_size > MAX_FILE_SIZE:
                error_msg = f"File size ({file_size / (1024 * 1024):.2f} MB) exceeds the maximum allowed size of 5 MB. Please upload a smaller file."
                yield f"data: {json.dumps({'status': 'ERROR', 'error': error_msg})}\n\n"
                return
            
            # Compute file hash for cache lookup
            file_hash = compute_file_hash(cv_content)
            print(f"CV file hash: {file_hash}")
            
            # Define async functions for parallel execution
            async def process_cv():
                """Process CV - Call #1"""
                # Check cache first
                cached_data = get_cv_cache(file_hash)
                
                if cached_data:
                    # Use cached data
                    print("Using cached CV extraction and fixed text")
                    cv_text = cached_data.get("raw_text", "")
                    fixed_text = cached_data.get("fixed_text", {})
                    return fixed_text, "CACHE_HIT"
                else:
                    # Extract text from file
                    cv_text = extract_text_from_file(cv_content, cv.filename)
                    
                    # Save raw extracted text (before LLM processing)
                    print(f"Saving raw extracted text (length: {len(cv_text)})")
                    
                    # Call LLM to fix/process the text
                    fixed_text = await fix_pdf_with_LayoutLM(cv_text)
                    
                    # Save to cache
                    save_cv_cache(file_hash, cv_text, fixed_text)
                    print("CV processing complete and cached")
                    return fixed_text, "PROCESSED"
            
            async def process_professor():
                """Process Professor Profile - Call #2+3"""
                # Check if professor exists in database (non-blocking)
                def check_professor_exists():
                    db = SessionLocal()
                    try:
                        return db.query(Professor).filter(Professor.root_url == root_url).first()
                    finally:
                        db.close()
                
                prof = await run_db_operation(check_professor_exists)
                
                # Get Professor Profile and Analysis
                professor_profile, professor_analysis_report = await get_professor_profile(root_url)
                return professor_profile, professor_analysis_report, prof is None
            
            # Start both tasks in parallel
            yield f"data: {json.dumps({'status': 'PARALLEL_PROCESSING', 'message': 'Processing CV and professor profile in parallel...'})}\n\n"
            
            # Send status updates
            yield f"data: {json.dumps({'status': 'PROCESSING_CV', 'message': 'Extracting and processing your CV...'})}\n\n"
            yield f"data: {json.dumps({'status': 'FETCHING_PROFILE', 'message': 'Fetching professor profile from database...'})}\n\n"
            
            # Create tasks for parallel execution
            cv_task = asyncio.create_task(process_cv())
            professor_task = asyncio.create_task(process_professor())
            
            # Wait for both to complete
            (fixed_text, cv_status), (professor_profile, professor_analysis_report, is_new_prof) = await asyncio.gather(
                cv_task,
                professor_task
            )
            
            # Send completion status updates
            if cv_status == "CACHE_HIT":
                yield f"data: {json.dumps({'status': 'CACHE_HIT', 'message': 'Using cached CV data...'})}\n\n"
            else:
                yield f"data: {json.dumps({'status': 'PROCESSING_CV_LLM', 'message': 'Processing CV with AI...'})}\n\n"
            
            if is_new_prof:
                yield f"data: {json.dumps({'status': 'CRAWLING', 'message': 'Professor not in database. Crawling website and performing detailed analysis...', 'info': 'This may take 3-5 minutes for a detailed analysis.'})}\n\n"
            else:
                yield f"data: {json.dumps({'status': 'ANALYZING', 'message': 'Analyzing professor profile and match...'})}\n\n"
            
            yield f"data: {json.dumps({'status': 'ANALYZING_PROFESSOR', 'message': 'Extracting and analyzing professor profile...'})}\n\n"
            
            # Step 3: Match Analysis & Refined Report Generation (Combined)
            yield f"data: {json.dumps({'status': 'MATCHING', 'message': 'Analyzing match and generating structured report...'})}\n\n"
            
            # Initialize variables
            match_analysis_report = None
            refined_report_json = None
            
            # Load combined prompt for match analysis and refined report generation
            combined_match_prompt = load_prompt(
                "match_analysis_and_refined_report.txt",
                student_profile=json.dumps(fixed_text, ensure_ascii=False),
                professor_profile=json.dumps(professor_profile, ensure_ascii=False),
                professor_report=professor_analysis_report
            )
            
            # Generate match analysis and refined report in one call
            student_name = fixed_text.get("identity", {}).get("full_name", "Unknown") if isinstance(fixed_text, dict) else "Unknown"
            prof_name = professor_profile.get("identity", {}).get("name", "Unknown")
            
            print("\n" + "="*60)
            print("🤖 LLM CALL #3: Combined Match Analysis & Refined Report Generation")
            print(f"   Model: gpt-5")
            print(f"   Purpose: Analyze match AND generate structured JSON report")
            print(f"   Student: {student_name}")
            print(f"   Professor: {prof_name}")
            print("="*60)
            
            try:
                async def make_match_analysis_call():
                    return await aclient.chat.completions.create(
                        model="gpt-5",
                        messages=[{"role": "user", "content": combined_match_prompt}],
                        response_format={"type": "json_object"},
                        stream=False
                    )
                
                combined_match_response = await retry_llm_call(make_match_analysis_call)
                
                combined_match_json_str = combined_match_response.choices[0].message.content
                print(f"\nLLM response received for combined match analysis + refined report. length: {len(combined_match_json_str)}")
                
                # Parse JSON response
                json_candidate = extract_json_by_brace_balance(combined_match_json_str)
                repaired = repair_json(json_candidate)
                combined_match_data = json.loads(repaired)
                
                # Extract match analysis report and refined report
                match_analysis_report = combined_match_data.get("match_analysis_report")
                refined_report_json = combined_match_data.get("refined_report")
                
                if not match_analysis_report:
                    raise Exception("Missing 'match_analysis_report' field in LLM response")
                if not refined_report_json:
                    raise Exception("Missing 'refined_report' field in LLM response")
                
                # Combine both reports for display
                # Format: Professor Analysis Report first, then Match Analysis Report
                combined_report = f"""# Professor & Lab Analysis Report

                    {professor_analysis_report}

                    ---

                    # Match Analysis Report

                    {match_analysis_report}
                    """
                
                # Save match log with combined report
                try:
                    save_match_log(combined_report, professor_profile, fixed_text)
                except Exception as log_error:
                    print(f"Warning: Failed to save match log: {log_error}")
                    # Continue even if log saving fails
                
                # Save refined report to metadata folder
                try:
                    save_refined_report(refined_report_json, professor_profile, fixed_text)
                except Exception as save_error:
                    print(f"Warning: Failed to save refined report: {save_error}")
                    # Continue even if saving fails
                
                # Save refined report to database (non-blocking)
                try:
                    await save_match_report_to_db(refined_report_json, professor_profile, fixed_text, root_url)
                except Exception as db_error:
                    print(f"Warning: Failed to save refined report to database: {db_error}")
                    # Continue even if database save fails
                
                print("Combined match analysis and refined report generated successfully")
                
            except Exception as match_error:
                print(f"Warning: Failed to generate combined match analysis and refined report: {match_error}")
                import traceback
                traceback.print_exc()
                # Set defaults if generation fails
                if match_analysis_report is None:
                    match_analysis_report = "Error: Failed to generate match analysis report."
                if refined_report_json is None:
                    refined_report_json = None
                
            # Create combined report for final output (even if generation failed)
            combined_report = f"""# Professor & Lab Analysis Report

                    {professor_analysis_report or "Error: Failed to generate professor analysis report."}

                    ---

                    # Match Analysis Report

                    {match_analysis_report if match_analysis_report else "Error: Failed to generate match analysis report."}
                    """
            
            # Send final result with both reports
            final_data = {
                'status': 'COMPLETE',
                'message': 'Analysis complete',
                'result': combined_report,
                'professor_report': professor_analysis_report,
                'match_report': match_analysis_report,
                'refined_report': refined_report_json
            }
            final_message = f"data: {json.dumps(final_data)}\n\n"
            print(f"DEBUG: Sending COMPLETE message, professor_report length: {len(professor_analysis_report or '')}, match_report length: {len(match_analysis_report or '')}")
            yield final_message
            print("DEBUG: COMPLETE message sent")
            
        except Exception as e:
            print(f"Error in analyze_match: {str(e)}")
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'status': 'ERROR', 'error': str(e)})}\n\n"
    
    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ==================== Async Database Helper Functions ====================

async def run_db_operation(func, *args, **kwargs):
    """
    Run a blocking database operation in a thread pool to avoid blocking the event loop
    """
    return await asyncio.to_thread(func, *args, **kwargs)

# ==================== Database Operations API ====================

@app.get("/dbp/professors")
async def list_professors(limit: Optional[int] = Query(None), offset: int = Query(0)):
    """List all professors"""
    try:
        # Run blocking DB operations in thread pool
        professors = await run_db_operation(db_ops.list_all, limit=limit, offset=offset)
        total_count = await run_db_operation(db_ops.count)
        
        # Convert to dict (this is CPU-bound, can also run in thread if needed)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors),
            "total": total_count
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/{professor_id}")
async def get_professor(professor_id: int):
    """Get professor by ID"""
    try:
        professor = await run_db_operation(db_ops.get_by_id, professor_id)
        if not professor:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Professor not found"}
            )
        professor_dict = await run_db_operation(db_ops.professor_to_dict, professor)
        return JSONResponse(content={
            "success": True,
            "data": professor_dict
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/url/{root_url:path}")
async def get_professor_by_url(root_url: str):
    """Get professor by root URL"""
    try:
        professor = await run_db_operation(db_ops.get_by_url, root_url)
        if not professor:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Professor not found"}
            )
        professor_dict = await run_db_operation(db_ops.professor_to_dict, professor)
        return JSONResponse(content={
            "success": True,
            "data": professor_dict
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/search")
async def search_professors_api(query: str = Query(...)):
    """Search professors"""
    try:
        professors = await run_db_operation(db_ops.search_professors, query)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/by-name")
async def get_professors_by_name(name: str = Query(...), exact_match: bool = Query(False)):
    """Get professors by name"""
    try:
        professors = await run_db_operation(db_ops.get_by_name, name, exact_match=exact_match)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/by-university")
async def get_professors_by_university(university: str = Query(...), exact_match: bool = Query(False)):
    """Get professors by university"""
    try:
        professors = await run_db_operation(db_ops.get_by_university, university, exact_match=exact_match)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/by-research")
async def get_professors_by_research(keyword: str = Query(...)):
    """Get professors by research interest"""
    try:
        professors = await run_db_operation(db_ops.get_professors_by_research_interest, keyword)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/professors/by-venue")
async def get_professors_by_venue(venue: str = Query(...)):
    """Get professors by venue"""
    try:
        professors = await run_db_operation(db_ops.get_professors_by_venue, venue)
        professors_data = [db_ops.professor_to_dict(p) for p in professors]
        return JSONResponse(content={
            "success": True,
            "data": professors_data,
            "count": len(professors)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.post("/dbp/professors")
async def create_professor_api(data: Dict[str, Any]):
    """Create a new professor"""
    try:
        required_fields = ["root_url", "name", "university", "profile_data"]
        for field in required_fields:
            if field not in data:
                return JSONResponse(
                    status_code=400,
                    content={"success": False, "error": f"Missing required field: {field}"}
                )
        
        professor = await run_db_operation(
            db_ops.create_professor,
            root_url=data["root_url"],
            name=data["name"],
            university=data["university"],
            profile_data=data["profile_data"]
        )
        professor_dict = await run_db_operation(db_ops.professor_to_dict, professor)
        return JSONResponse(content={
            "success": True,
            "data": professor_dict
        })
    except ValueError as e:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": str(e)}
        )
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.put("/dbp/professors/{professor_id}")
async def update_professor_api(professor_id: int, data: Dict[str, Any]):
    """Update professor by ID"""
    try:
        professor = await run_db_operation(
            db_ops.update_professor,
            professor_id=professor_id,
            name=data.get("name"),
            university=data.get("university"),
            profile_data=data.get("profile_data")
        )
        if not professor:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Professor not found"}
            )
        professor_dict = await run_db_operation(db_ops.professor_to_dict, professor)
        return JSONResponse(content={
            "success": True,
            "data": professor_dict
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.delete("/dbp/professors/{professor_id}")
async def delete_professor_api(professor_id: int):
    """Delete professor by ID"""
    try:
        success = await run_db_operation(db_ops.delete_by_id, professor_id)
        if not success:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Professor not found"}
            )
        return JSONResponse(content={
            "success": True,
            "message": "Professor deleted successfully"
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.delete("/dbp/professors/url/{root_url:path}")
async def delete_professor_by_url_api(root_url: str):
    """Delete professor by root URL"""
    try:
        success = await run_db_operation(db_ops.delete_by_url, root_url)
        if not success:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Professor not found"}
            )
        return JSONResponse(content={
            "success": True,
            "message": "Professor deleted successfully"
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/dbp/stats")
async def get_stats():
    """Get database statistics"""
    try:
        total_count = await run_db_operation(db_ops.count)
        
        # Run match reports count in thread pool
        def count_match_reports():
            db = SessionLocal()
            try:
                return db.query(MatchReport).count()
            finally:
                db.close()
        
        match_reports_count = await run_db_operation(count_match_reports)
        
        return JSONResponse(content={
            "success": True,
            "data": {
                "total_professors": total_count,
                "total_match_reports": match_reports_count
            }
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


# ==================== Match Reports API ====================

@app.get("/api/match-reports")
async def get_all_match_reports(limit: Optional[int] = Query(50), offset: int = Query(0)):
    """Get all match reports (for history)"""
    try:
        # Run blocking DB query in thread pool
        def fetch_match_reports():
            db = SessionLocal()
            try:
                reports = db.query(MatchReport).order_by(MatchReport.id.desc()).offset(offset).limit(limit).all()
                reports_list = []
                for report in reports:
                    reports_list.append({
                        "id": report.id,
                        "professor_name": report.professor_name,
                        "student_name": report.student_name,
                        "professor_url": report.professor_url,
                        "created_at": report.created_at,
                        "overall_score": report.overall_score
                    })
                return reports_list
            finally:
                db.close()
        
        reports_list = await run_db_operation(fetch_match_reports)
        
        return JSONResponse(content={
            "success": True,
            "data": reports_list,
            "count": len(reports_list)
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


@app.get("/api/match-reports/{report_id}")
async def get_match_report(report_id: int):
    """Get a specific match report by ID"""
    try:
        # Run blocking DB query in thread pool
        def fetch_match_report():
            db = SessionLocal()
            try:
                report = db.query(MatchReport).filter(MatchReport.id == report_id).first()
                if not report:
                    return None
                return {
                    "id": report.id,
                    "professor_name": report.professor_name,
                    "student_name": report.student_name,
                    "professor_url": report.professor_url,
                    "created_at": report.created_at,
                    "overall_score": report.overall_score,
                    "refined_report": report.refined_report
                }
            finally:
                db.close()
        
        report_data = await run_db_operation(fetch_match_report)
        
        if not report_data:
            return JSONResponse(
                status_code=404,
                content={"success": False, "error": "Match report not found"}
            )
        
        return JSONResponse(content={
            "success": True,
            "data": report_data
        })
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": str(e)}
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)